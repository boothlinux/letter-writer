#Copyright 2021 Colton Booth
#Released under the Apache 2.0 licence
#See LICENCE file for complete text

from docx import Document
import PySimpleGUI as sg
from datetime import datetime
import time

sg.theme('SystemDefault')
layout = [[sg.Text("Name:")],
        [sg.Input(key='name')],
        [sg.Text("Address:")],
        [sg.MLine(default_text='', size=(45, 3),
                  key='address')],
        [sg.Text("Salutation:")],
        [sg.Input(default_text='Dear ', key='salutation')],
        [sg.Text("Body Text")],
        [sg.MLine(default_text='', size=(45, 3),
                  key='bodytext')],
        [sg.Text("Signature")],
        [sg.MLine(default_text='Regards, ', size=(45, 3),
                  key='signature')],
        [sg.Text(size=(40,1), key='outputline')],
        [sg.Button('Ok'), sg.Button('Quit')]]

window = sg.Window('Letter Writer v0.2', layout)

while True:
    event, values = window.read()
    if event == sg.WINDOW_CLOSED or event == 'Quit':
        break

    try:
        #extract the variables
        name = values["name"]
        address = values["address"]
        salutation = values["salutation"]
        bodytext = values["bodytext"]
        signature = values["signature"]
        now = datetime.utcnow()
        todays_date = now.strftime('%Y-%m-%d')

        #prepare the letter
        try:
            document = Document('letterhead.docx')
        except:
            document = Document()

        p = document.add_paragraph(todays_date)
        p = document.add_paragraph(" ")
        p = document.add_paragraph(" ")
        p = document.add_paragraph(name + '\n'  + address )
        p = document.add_paragraph(" ")
        p = document.add_paragraph(" ")
        p = document.add_paragraph(salutation + " " + name + ",")
        p = document.add_paragraph(" ")
        p = document.add_paragraph(bodytext)
        p = document.add_paragraph(" ")
        p = document.add_paragraph(signature)

        document.save(name + " - " + todays_date +".docx")
        window['outputline'].update('Created!')

    except Exception as e:
        print(e)
        window['outputline'].update('ERROR! Something is wrong.')

window.close()
